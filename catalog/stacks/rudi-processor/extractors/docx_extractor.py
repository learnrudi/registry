#!/usr/bin/env python3
"""
DOCX Extractor - Extracts text from Word documents
"""

from .base_extractor import BaseExtractor
from typing import Dict, Any

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False


class DocxExtractor(BaseExtractor):
    """Extract content from Office documents"""

    def extract(self, file_path: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Extract text from DOCX/XLSX files"""

        extension = metadata['basic_metadata']['extension']

        if extension == '.docx':
            if not HAS_DOCX:
                metadata['processing_status']['errors'].append("python-docx not available")
                return metadata

            try:
                doc = Document(file_path)

                # Extract all text from paragraphs
                full_text = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)

                # Extract text from tables
                table_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            table_text.append(' | '.join(row_text))

                if table_text:
                    full_text.append('\n--- Tables ---\n')
                    full_text.extend(table_text)

                # Store full text
                metadata['extracted_content']['full_text'] = '\n'.join(full_text)

                # Count elements
                metadata['extracted_content']['paragraph_count'] = len(doc.paragraphs)
                metadata['extracted_content']['table_count'] = len(doc.tables)
                metadata['extracted_content']['word_count'] = len(' '.join(full_text).split())

                # Extract core properties if available
                if doc.core_properties:
                    props = {}
                    if doc.core_properties.title:
                        props['title'] = doc.core_properties.title
                    if doc.core_properties.author:
                        props['author'] = doc.core_properties.author
                    if doc.core_properties.subject:
                        props['subject'] = doc.core_properties.subject
                    if doc.core_properties.created:
                        props['created'] = str(doc.core_properties.created)
                    if doc.core_properties.modified:
                        props['modified'] = str(doc.core_properties.modified)

                    if props:
                        metadata['extracted_content']['document_properties'] = props

            except Exception as e:
                metadata['processing_status']['errors'].append(f"DOCX extraction error: {str(e)}")

        elif extension == '.xlsx':
            if not HAS_EXCEL:
                metadata['processing_status']['errors'].append("openpyxl not available")
                return metadata

            try:
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)

                # Extract text from all sheets
                full_text = []
                sheet_info = []

                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    sheet_text = [f"--- Sheet: {sheet_name} ---"]

                    row_count = 0
                    for row in sheet.iter_rows(values_only=True):
                        # Filter out empty rows
                        row_values = [str(cell) if cell is not None else '' for cell in row]
                        if any(v.strip() for v in row_values):
                            sheet_text.append('\t'.join(row_values))
                            row_count += 1

                    if row_count > 0:
                        full_text.extend(sheet_text)
                        sheet_info.append({
                            'name': sheet_name,
                            'rows': row_count
                        })

                metadata['extracted_content']['full_text'] = '\n'.join(full_text)
                metadata['extracted_content']['sheet_count'] = len(wb.sheetnames)
                metadata['extracted_content']['sheets'] = sheet_info

                wb.close()

            except Exception as e:
                metadata['processing_status']['errors'].append(f"XLSX extraction error: {str(e)}")

        else:
            metadata['processing_status']['errors'].append(f"Unsupported Office format: {extension}")

        return metadata
